# Crear archivo Word con guía y preguntas frecuentes
from docx import Document

doc = Document()

# Título principal
doc.add_heading('Guía/Esquema para Explicación del Proyecto MyDeptos', 0)

# Sección A
doc.add_heading('A. Introducción', level=1)
doc.add_paragraph('Nombre del proyecto: MyDeptos')
doc.add_paragraph('Objetivo: Automatizar y facilitar la búsqueda, publicación y gestión de departamentos para alquiler o venta, centralizando la experiencia tanto para usuarios comunes como para propietarios/inmobiliarias.')

# Sección B
doc.add_heading('B. Estructura General', level=1)
doc.add_paragraph('Tecnologías principales:')
doc.add_paragraph('- Backend: Flask (Python)\n- Base de datos: MySQL\n- Frontend: HTML, CSS, Bootstrap, Jinja2\n- Otros: Flask-WTF (formularios), Flask-Login (autenticación), Bcrypt (seguridad), AJAX (favoritos, notificaciones)')

doc.add_paragraph('Principales módulos/rutas:')
doc.add_paragraph('- Autenticación: Registro, login, logout, recuperación y cambio de contraseña.\n- Panel de usuario: Visualización y gestión de publicaciones, edición de perfil, notificaciones.\n- Publicación de departamentos: Formulario con validaciones, subida de imágenes, ubicación en mapa.\n- Búsqueda y filtrado: Por tipo, precio, localidad, ambientes, etc.\n- Favoritos: Añadir/quitar favoritos, listado personalizado.\n- Reseñas: Dejar, editar y eliminar reseñas sobre departamentos.\n- Notificaciones: Sistema de avisos para eventos importantes (nuevas reseñas, etc.).\n- Gestión de publicaciones: Modificar/eliminar departamentos propios.')

# Sección C
doc.add_heading('C. Flujo de Usuario', level=1)
doc.add_paragraph('1. Registro e inicio de sesión\n2. Publicación de un departamento (con fotos y ubicación)\n3. Búsqueda y filtrado de departamentos\n4. Agregar a favoritos\n5. Dejar una reseña\n6. Gestión de notificaciones\n7. Edición de perfil y publicaciones')

# Sección D
doc.add_heading('D. Seguridad y buenas prácticas', level=1)
doc.add_paragraph('- Hash de contraseñas (bcrypt/pbkdf2)\n- CSRF en formularios\n- Validación de datos y archivos\n- Manejo de sesiones y roles')

# Sección E
doc.add_heading('E. Mejoras recomendadas para excelencia', level=1)
doc.add_paragraph('- Pruebas unitarias y de integración\n- Mejor documentación (README, docstrings)\n- Manejo avanzado de errores (páginas personalizadas)\n- Optimización de consultas SQL\n- Internacionalización (multi-idioma)\n- Mejor UX/UI (más AJAX, feedback visual)\n- Despliegue en un servidor real (Heroku, PythonAnywhere, etc.)')

# Preguntas frecuentes
doc.add_heading('Preguntas Frecuentes para el Tribunal', 0)

doc.add_heading('Sobre funcionalidad y uso', level=1)
preguntas_funcionalidad = [
    "¿Cuál es el objetivo principal de MyDeptos y qué problema resuelve?",
    "¿Cómo se diferencia tu sistema de otras plataformas de búsqueda de departamentos?",
    "¿Qué roles de usuario existen y qué puede hacer cada uno?",
    "¿Cómo es el flujo típico de un usuario desde el registro hasta la publicación y gestión de un departamento?",
    "¿Cómo se gestionan los favoritos y las reseñas?",
    "¿Qué sucede si un usuario olvida su contraseña?",
    "¿Cómo se notifican los eventos importantes a los usuarios?",
    "¿Qué validaciones existen al publicar un departamento?",
    "¿Cómo se asegura la calidad de las imágenes subidas?"
]
for p in preguntas_funcionalidad:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Preguntas técnicas', level=1)
preguntas_tecnicas = [
    "¿Por qué elegiste Flask y MySQL para el desarrollo?",
    "¿Cómo manejas la seguridad de las contraseñas?",
    "¿Cómo implementaste la protección CSRF en los formularios?",
    "¿Cómo gestionas las sesiones de usuario y la autenticación?",
    "¿Cómo está estructurada la base de datos?",
    "¿Cómo se realiza la búsqueda y filtrado de departamentos?",
    "¿Cómo gestionas los errores y los logs en la aplicación?",
    "¿Cómo se almacenan y sirven las imágenes de los departamentos?",
    "¿Cómo implementaste el sistema de notificaciones?",
    "¿Qué ocurre si hay un error al subir una imagen o guardar un departamento?"
]
for p in preguntas_tecnicas:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Preguntas de mejora y escalabilidad', level=1)
preguntas_mejora = [
    "¿Qué mejoras implementarías si tuvieras más tiempo?",
    "¿Cómo escalarías la aplicación para soportar más usuarios?",
    "¿Cómo agregarías soporte para otros idiomas?",
    "¿Cómo implementarías pruebas automáticas?",
    "¿Cómo desplegarías la aplicación en un entorno de producción?",
    "¿Qué harías para mejorar la seguridad?"
]
for p in preguntas_mejora:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Preguntas de justificación y reflexión', level=1)
preguntas_reflexion = [
    "¿Qué fue lo más desafiante del desarrollo?",
    "¿Qué aprendiste durante el proyecto?",
    "¿Por qué crees que tu solución es útil y aporta valor?",
    "¿Cómo validaste que tu sistema cumple con los objetivos planteados?",
    "¿Qué feedback recibiste de usuarios de prueba (si hubo)?"
]
for p in preguntas_reflexion:
    doc.add_paragraph(p, style='List Bullet')

# Guardar el archivo
doc.save('Guia_MyDeptos.docx')
print("Archivo 'Guia_MyDeptos.docx' creado correctamente.")